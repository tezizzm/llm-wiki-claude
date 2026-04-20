import importlib.util
import json
import sys
from datetime import datetime as real_datetime
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

INGEST_PATH = ROOT / "scripts" / "ingest.py"
spec = importlib.util.spec_from_file_location("ingest_module", INGEST_PATH)
ingest = importlib.util.module_from_spec(spec)
assert spec and spec.loader
spec.loader.exec_module(ingest)

from scripts.workspace import resolve_workspace


def _make_workspace(root: Path):
    """Return a WorkspacePaths rooted at ``root`` with source='flag'.

    Tests call this instead of the ``tmp_workspace`` fixture when they need a
    bespoke layout (e.g. seeding schemas/, an .ingest-settings.json at the
    workspace root, or a pre-populated raw/inbox before invoking ingest.main).
    """
    root.mkdir(parents=True, exist_ok=True)
    return resolve_workspace(str(root), None)


def test_template_ingest_settings_is_shareable():
    config = json.loads(Path("ingest-settings.json").read_text(encoding="utf-8"))
    assert config["schema_version"] == 1
    assert config["max_topics"] == 6
    assert config["max_entities"] == 6
    assert "opaque_task_regex" in config["low_signal_sources"]


def test_load_ingest_settings_compiles_low_signal_regex(tmp_path):
    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir()
    (workspace_root / "ingest-settings.local.json").write_text(
        json.dumps({"max_topics": 5}), encoding="utf-8"
    )
    workspace = _make_workspace(workspace_root)

    settings = ingest.load_ingest_settings(workspace)

    assert settings["max_topics"] == 5
    assert settings["low_signal_sources"]["opaque_task_regex_compiled"].match("AM-ab12.md")


def test_prepare_ingest_settings_warns_for_deprecated_config_version(tmp_path):
    settings_path = tmp_path / "ingest-settings.json"
    prepared, warnings = ingest.prepare_ingest_settings({"config_version": 1}, settings_path)

    assert prepared["schema_version"] == 1
    assert "deprecated" in warnings[0]


def test_is_low_signal_source_uses_settings():
    settings = ingest.load_ingest_settings(resolve_workspace(None, None))
    assert ingest.is_low_signal_source(Path("AM-ab12.md"), settings) is True
    assert ingest.is_low_signal_source(Path("product_spec.md"), settings) is False


def test_select_terms_filters_topics_and_entities():
    settings = ingest.load_ingest_settings(resolve_workspace(None, None))

    topics = ingest.select_terms(
        ["Approval Workflow", "Issue21", "pkg/runtime/provider.go", "Approval Workflow"],
        settings["max_topics"],
        ingest.is_low_value_topic,
        settings,
    )
    entities = ingest.select_terms(
        ["AgentMesh", "ExecDispatcher", "Kubernetes", "future stakeholder"],
        settings["max_entities"],
        ingest.is_low_value_entity,
        settings,
    )

    assert topics == ["Approval Workflow"]
    assert entities == ["AgentMesh", "Kubernetes"]


def test_frontmatter_value_reads_title():
    text = "---\ntitle: \"Hello World\"\nstatus: draft\n---\nbody"
    assert ingest.frontmatter_value(text, "title") == "Hello World"
    assert ingest.frontmatter_value(text, "missing") is None


def test_init_client_requires_real_api_key(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "your_anthropic_api_key_here")
    try:
        ingest.init_client()
    except RuntimeError as exc:
        assert "Missing ANTHROPIC_API_KEY" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for placeholder API key")


def test_save_last_ingest_run_writes_summary(tmp_path):
    workspace = _make_workspace(tmp_path / "workspace")

    ingest.save_last_ingest_run(workspace, {"status": "completed", "processed": 2})

    saved = json.loads(workspace.last_ingest_run_path.read_text(encoding="utf-8"))
    assert saved["status"] == "completed"
    assert saved["processed"] == 2


def test_upsert_source_contribution_replaces_existing_block(tmp_path):
    target = tmp_path / "topic.md"
    target.write_text(
        ingest.render_contribution_block("old.md", "# Old\n\nOld content") + "\n",
        encoding="utf-8",
    )

    ingest.upsert_source_contribution(target, "old.md", "# Old\n\nNew content")

    content = target.read_text(encoding="utf-8")
    assert content.count("<!-- SOURCE:old.md -->") == 1
    assert "New content" in content
    assert "Old content" not in content


def test_remove_source_contribution_deletes_empty_file(tmp_path):
    target = tmp_path / "entity.md"
    target.write_text(
        ingest.render_contribution_block("demo.md", "# Entity\n\nOnly content") + "\n",
        encoding="utf-8",
    )

    removed = ingest.remove_source_contribution(target, "demo.md")

    assert removed is True
    assert not target.exists()


def test_cleanup_source_artifacts_uses_manifest_record(tmp_path):
    workspace = _make_workspace(tmp_path / "workspace")
    summaries = workspace.summaries_dir
    topics = workspace.topics_dir
    entities = workspace.entities_dir
    summaries.mkdir(parents=True)
    topics.mkdir(parents=True)
    entities.mkdir(parents=True)

    (summaries / "demo.md").write_text("# Summary", encoding="utf-8")
    ingest.upsert_source_contribution(topics / "agent-router.md", "demo.md", "# Agent Router")
    ingest.upsert_source_contribution(entities / "demomesh.md", "demo.md", "# DemoMesh")

    removed = ingest.cleanup_source_artifacts(
        workspace,
        "demo.md",
        {
            "summary_path": "wiki/summaries/demo.md",
            "topic_slugs": ["agent-router"],
            "entity_slugs": ["demomesh"],
        },
    )

    assert removed == {"summaries": 1, "topics": 1, "entities": 1}
    assert not (summaries / "demo.md").exists()
    assert not (topics / "agent-router.md").exists()
    assert not (entities / "demomesh.md").exists()


def test_refine_merged_pages_merges_plural_alias_and_normalizes(tmp_path):
    workspace = _make_workspace(tmp_path / "workspace")
    topics = workspace.topics_dir
    entities = workspace.entities_dir
    summaries = workspace.summaries_dir
    topics.mkdir(parents=True)
    entities.mkdir(parents=True)
    summaries.mkdir(parents=True)

    ingest.upsert_source_contribution(topics / "approval-workflows.md", "one.md", "# Approval Workflows\n\nOne")
    ingest.upsert_source_contribution(topics / "approval-workflow.md", "two.md", "# Approval Workflow\n\nTwo")

    actions = ingest.refine_merged_pages(
        workspace, ingest.load_ingest_settings(resolve_workspace(None, None))
    )

    assert actions["topic_alias_merges"] == 1
    merged = (topics / "approval-workflow.md").read_text(encoding="utf-8")
    assert "one.md" in merged
    assert "two.md" in merged
    assert not (topics / "approval-workflows.md").exists()


def test_write_last_ingest_report_writes_markdown_report(tmp_path):
    workspace = _make_workspace(tmp_path / "workspace")

    ingest.write_last_ingest_report(
        workspace,
        {
            "status": "completed",
            "ran_at_utc": "2026-01-02T03:04:05Z",
            "model": "fake-model",
            "raw_candidates": 3,
            "processed": 2,
            "reconciled": False,
            "cleaned_stale": {"manifest_entries": 1},
            "merge_cleanup": {"topic_alias_merges": 1, "entity_alias_merges": 0, "topic_low_value_prunes": 0, "entity_low_value_prunes": 0},
            "page_stats": {"summaries": 2, "topics": 3, "entities": 1, "topic_contributions": 4, "entity_contributions": 1},
            "processed_files": ["a.md", "b.md"],
            "errors": [],
        }
    )

    report = workspace.ingest_report_path.read_text(encoding="utf-8")
    assert "# Last Ingest Report" in report
    assert "`completed`" in report
    assert "`a.md`" in report
    assert "Topic alias merges" in report


def test_extract_text_supports_pdf_via_pdf_reader(monkeypatch, tmp_path):
    pdf_path = tmp_path / "demo.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    class FakePage:
        def extract_text(self):
            return "PDF content"

    class FakeReader:
        def __init__(self, path):
            self.pages = [FakePage()]

    monkeypatch.setattr(ingest, "PdfReader", FakeReader)
    assert ingest.extract_text(pdf_path) == "PDF content"


def test_ingest_main_builds_expected_summary_from_fixture(tmp_path, monkeypatch):
    fixture_root = Path("tests/fixtures/ingest/basic")
    raw_fixture = fixture_root / "raw" / "demo.md"
    expected_summary = (fixture_root / "golden" / "summary.md").read_text(encoding="utf-8")

    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir()
    raw_dir = workspace_root / "raw" / "inbox"
    schemas_dir = workspace_root / "schemas"
    raw_dir.mkdir(parents=True)
    schemas_dir.mkdir(parents=True)
    (schemas_dir / "AGENTS.md").write_text("Test schema", encoding="utf-8")
    (workspace_root / ".wikiignore").write_text("", encoding="utf-8")
    (workspace_root / "ingest-settings.json").write_text(
        Path("ingest-settings.json").read_text(encoding="utf-8"), encoding="utf-8"
    )
    (raw_dir / "demo.md").write_text(raw_fixture.read_text(encoding="utf-8"), encoding="utf-8")

    workspace = _make_workspace(workspace_root)
    state_dir = workspace.state_dir
    wiki_dir = workspace.wiki_dir

    class FixedDateTime:
        @staticmethod
        def now(tz=None):
            return real_datetime(2026, 1, 2, 3, 4, 5)

    monkeypatch.setattr(ingest, "datetime", FixedDateTime)
    monkeypatch.setattr(ingest, "init_client", lambda: ("fake-client", "fake-model"))
    monkeypatch.setattr(
        ingest,
        "call_claude_json",
        lambda client, model, system_prompt, user_prompt: {
            "title": "Demo Knowledge Page",
            "summary": "AgentMesh coordinates distributed agent work through a registry and router.",
            "key_facts": [
                "AgentMesh has a capability registry",
                "AgentMesh has an agent router",
            ],
            "topics": ["Capability Registry", "Agent Router"],
            "entities": ["AgentMesh", "Dapr"],
            "open_questions": ["How should routing policies evolve?"],
            "topic_summaries": {
                "Capability Registry": "Registry for agent capabilities.",
                "Agent Router": "Routes work to the right capability.",
            },
            "entity_summaries": {
                "AgentMesh": "Distributed control plane for agent ecosystems.",
                "Dapr": "Runtime substrate used by AgentMesh.",
            },
        },
    )

    rc = ingest.main([], workspace)
    assert rc == 0

    summary_path = wiki_dir / "summaries" / "demo.md"
    assert summary_path.read_text(encoding="utf-8") == expected_summary
    manifest = json.loads((state_dir / "manifest.json").read_text(encoding="utf-8"))
    assert "raw/inbox/demo.md" in manifest["files"]
    assert manifest["files"]["raw/inbox/demo.md"]["topic_slugs"] == ["agent-router", "capability-registry"]
    run_summary = json.loads((state_dir / "last_ingest_run.json").read_text(encoding="utf-8"))
    assert run_summary["processed"] == 1
    assert run_summary["page_stats"]["topics"] == 2
    event_lines = (state_dir / "ingest_events.jsonl").read_text(encoding="utf-8").strip().splitlines()
    assert len(event_lines) == 2
    assert json.loads(event_lines[0])["event"] == "ingest_file_started"
    assert json.loads(event_lines[1])["event"] == "ingest_file_completed"
    assert (state_dir / "last_ingest_report.md").exists()


def test_ingest_dry_run_reports_actions_without_writing(tmp_path):
    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir()
    raw_dir = workspace_root / "raw" / "inbox"
    raw_dir.mkdir(parents=True)
    (raw_dir / "demo.md").write_text("hello", encoding="utf-8")
    (workspace_root / ".wikiignore").write_text("", encoding="utf-8")
    (workspace_root / "schemas").mkdir()
    (workspace_root / "schemas" / "AGENTS.md").write_text("schema", encoding="utf-8")
    (workspace_root / "ingest-settings.json").write_text(
        Path("ingest-settings.json").read_text(encoding="utf-8"), encoding="utf-8"
    )

    workspace = _make_workspace(workspace_root)
    rc = ingest.main(["--dry-run"], workspace)
    assert rc == 0

    run_summary = json.loads(workspace.last_ingest_run_path.read_text(encoding="utf-8"))
    assert run_summary["status"] == "dry_run"
    assert run_summary["would_process"] == ["demo.md"]
    # Dry-run still goes through ensure_workspace_writable() per the DISPATCH
    # contract (ARCHITECTURE §5.3 / §6), so wiki subdirs exist but are empty.
    assert not list((workspace.summaries_dir).glob("*"))
    assert not workspace.ingest_events_path.exists()
    assert workspace.ingest_report_path.exists()


def test_ingest_reconcile_resets_stale_outputs(tmp_path, monkeypatch):
    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir()
    raw_dir = workspace_root / "raw" / "inbox"
    wiki_dir = workspace_root / "wiki"
    state_dir = workspace_root / "state"
    schemas_dir = workspace_root / "schemas"
    raw_dir.mkdir(parents=True)
    (wiki_dir / "summaries").mkdir(parents=True)
    state_dir.mkdir(parents=True)
    schemas_dir.mkdir(parents=True)
    (raw_dir / "fresh.md").write_text("fresh", encoding="utf-8")
    (wiki_dir / "summaries" / "stale.md").write_text("old", encoding="utf-8")
    (state_dir / "manifest.json").write_text(
        json.dumps({"files": {"raw/inbox/stale.md": {"sha256": "old"}}}),
        encoding="utf-8",
    )
    (workspace_root / ".wikiignore").write_text("", encoding="utf-8")
    (schemas_dir / "AGENTS.md").write_text("schema", encoding="utf-8")
    (workspace_root / "ingest-settings.json").write_text(
        Path("ingest-settings.json").read_text(encoding="utf-8"), encoding="utf-8"
    )

    monkeypatch.setattr(ingest, "init_client", lambda: ("fake-client", "fake-model"))
    monkeypatch.setattr(
        ingest,
        "call_claude_json",
        lambda *args, **kwargs: {
            "title": "Fresh",
            "summary": "Fresh summary",
            "key_facts": [],
            "topics": [],
            "entities": [],
            "open_questions": [],
            "topic_summaries": {},
            "entity_summaries": {},
        },
    )

    workspace = _make_workspace(workspace_root)
    rc = ingest.main(["--reconcile"], workspace)
    assert rc == 0

    assert not (wiki_dir / "summaries" / "stale.md").exists()
    assert (wiki_dir / "summaries" / "fresh.md").exists()


def test_ingest_cleans_stale_manifest_entries_without_full_reconcile(tmp_path, monkeypatch):
    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir()
    raw_dir = workspace_root / "raw" / "inbox"
    wiki_dir = workspace_root / "wiki"
    state_dir = workspace_root / "state"
    schemas_dir = workspace_root / "schemas"
    raw_dir.mkdir(parents=True)
    (wiki_dir / "summaries").mkdir(parents=True)
    (wiki_dir / "topics").mkdir(parents=True)
    state_dir.mkdir(parents=True)
    schemas_dir.mkdir(parents=True)
    (raw_dir / "fresh.md").write_text("fresh", encoding="utf-8")
    (wiki_dir / "summaries" / "stale.md").write_text("old", encoding="utf-8")
    ingest.upsert_source_contribution(wiki_dir / "topics" / "legacy-topic.md", "stale.md", "# Legacy")
    (state_dir / "manifest.json").write_text(
        json.dumps(
            {
                "files": {
                    "raw/inbox/stale.md": {
                        "sha256": "old",
                        "summary_path": "wiki/summaries/stale.md",
                        "topic_slugs": ["legacy-topic"],
                        "entity_slugs": [],
                    }
                }
            }
        ),
        encoding="utf-8",
    )
    (workspace_root / ".wikiignore").write_text("", encoding="utf-8")
    (schemas_dir / "AGENTS.md").write_text("schema", encoding="utf-8")
    (workspace_root / "ingest-settings.json").write_text(
        Path("ingest-settings.json").read_text(encoding="utf-8"), encoding="utf-8"
    )

    monkeypatch.setattr(ingest, "init_client", lambda: ("fake-client", "fake-model"))
    monkeypatch.setattr(
        ingest,
        "call_claude_json",
        lambda *args, **kwargs: {
            "title": "Fresh",
            "summary": "Fresh summary",
            "key_facts": [],
            "topics": [],
            "entities": [],
            "open_questions": [],
            "topic_summaries": {},
            "entity_summaries": {},
        },
    )

    workspace = _make_workspace(workspace_root)
    rc = ingest.main([], workspace)
    assert rc == 0

    run_summary = json.loads((state_dir / "last_ingest_run.json").read_text(encoding="utf-8"))
    assert run_summary["reconciled"] is False
    assert run_summary["cleaned_stale"]["manifest_entries"] == 1
    assert not (wiki_dir / "summaries" / "stale.md").exists()
    assert not (wiki_dir / "topics" / "legacy-topic.md").exists()


def test_ingest_updates_source_contributions_when_topics_change(tmp_path, monkeypatch):
    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir()
    raw_dir = workspace_root / "raw" / "inbox"
    wiki_dir = workspace_root / "wiki"
    state_dir = workspace_root / "state"
    schemas_dir = workspace_root / "schemas"
    raw_dir.mkdir(parents=True)
    (wiki_dir / "topics").mkdir(parents=True)
    state_dir.mkdir(parents=True)
    schemas_dir.mkdir(parents=True)
    (raw_dir / "demo.md").write_text("demo", encoding="utf-8")
    (workspace_root / ".wikiignore").write_text("", encoding="utf-8")
    (schemas_dir / "AGENTS.md").write_text("schema", encoding="utf-8")
    (workspace_root / "ingest-settings.json").write_text(
        Path("ingest-settings.json").read_text(encoding="utf-8"), encoding="utf-8"
    )

    monkeypatch.setattr(ingest, "get_schema_text", lambda ws: "schema")

    responses = iter(
        [
            {
                "title": "Demo",
                "summary": "First pass",
                "key_facts": [],
                "topics": ["Alpha Topic"],
                "entities": [],
                "open_questions": [],
                "topic_summaries": {"Alpha Topic": "Alpha"},
                "entity_summaries": {},
            },
            {
                "title": "Demo",
                "summary": "Second pass",
                "key_facts": [],
                "topics": ["Beta Topic"],
                "entities": [],
                "open_questions": [],
                "topic_summaries": {"Beta Topic": "Beta"},
                "entity_summaries": {},
            },
        ]
    )
    monkeypatch.setattr(ingest, "call_claude_json", lambda *args, **kwargs: next(responses))

    workspace = _make_workspace(workspace_root)
    settings = ingest.load_ingest_settings(workspace)

    result_one = ingest.ingest_file(workspace, "fake-client", "fake-model", raw_dir / "demo.md", settings)
    result_two = ingest.ingest_file(
        workspace,
        "fake-client",
        "fake-model",
        raw_dir / "demo.md",
        settings,
        previous_record=result_one,
    )

    assert result_two["topic_slugs"] == ["beta-topic"]
    assert not (wiki_dir / "topics" / "alpha-topic.md").exists()
    beta_content = (wiki_dir / "topics" / "beta-topic.md").read_text(encoding="utf-8")
    assert "<!-- SOURCE:demo.md -->" in beta_content


# --- HTML extraction tests (LWC-wyli) ---


def test_extract_html_text_strips_tags_and_returns_text(tmp_path):
    html_file = tmp_path / "sample.html"
    html_file.write_text(
        "<html><body><h1>Title</h1><p>Hello <b>world</b></p></body></html>",
        encoding="utf-8",
    )
    result = ingest.extract_html_text(html_file)
    assert "Title" in result
    assert "Hello" in result
    assert "world" in result
    assert "<h1>" not in result
    assert "<p>" not in result
    assert "<b>" not in result


def test_extract_html_text_excludes_script_and_style(tmp_path):
    html_file = tmp_path / "scripted.html"
    html_file.write_text(
        "<html><head><style>body { color: red; }</style>"
        "<script>alert('hi');</script></head>"
        "<body><p>Visible text</p></body></html>",
        encoding="utf-8",
    )
    result = ingest.extract_html_text(html_file)
    assert "Visible text" in result
    assert "color: red" not in result
    assert "alert" not in result


def test_extract_html_text_empty_html_returns_placeholder(tmp_path):
    html_file = tmp_path / "empty.html"
    html_file.write_text("<html><body></body></html>", encoding="utf-8")
    result = ingest.extract_html_text(html_file)
    assert result == "[HTML parsed but no extractable text found: empty.html]"


def test_extract_html_text_tags_only_returns_placeholder(tmp_path):
    html_file = tmp_path / "tags_only.html"
    html_file.write_text(
        "<html><head><title></title></head><body><div><span></span></div></body></html>",
        encoding="utf-8",
    )
    result = ingest.extract_html_text(html_file)
    assert result == "[HTML parsed but no extractable text found: tags_only.html]"


def test_extract_text_dispatches_htm_extension(tmp_path):
    htm_file = tmp_path / "page.htm"
    htm_file.write_text(
        "<html><body><p>HTM content</p></body></html>",
        encoding="utf-8",
    )
    result = ingest.extract_text(htm_file)
    assert "HTM content" in result


def test_extract_text_dispatches_html_extension(tmp_path):
    html_file = tmp_path / "page.html"
    html_file.write_text(
        "<html><body><p>HTML content</p></body></html>",
        encoding="utf-8",
    )
    result = ingest.extract_text(html_file)
    assert "HTML content" in result


def test_extract_html_text_latin1_fallback(tmp_path):
    html_file = tmp_path / "latin1.html"
    # Write bytes that are valid Latin-1 but invalid UTF-8
    html_file.write_bytes(
        b"<html><body><p>caf\xe9</p></body></html>"
    )
    result = ingest.extract_html_text(html_file)
    assert "caf" in result


# --- RST extraction tests (LWC-8tav) ---


def test_extract_rst_text_converts_headings_and_bold(tmp_path):
    rst_file = tmp_path / "sample.rst"
    rst_file.write_text(
        "My Title\n"
        "========\n"
        "\n"
        "Some **bold** text here.\n"
        "\n"
        "* Item one\n"
        "* Item two\n",
        encoding="utf-8",
    )
    result = ingest.extract_rst_text(rst_file)
    assert "My Title" in result
    assert "bold" in result
    assert "Item one" in result
    assert "Item two" in result
    # RST directives should be stripped
    assert "========" not in result
    assert "**" not in result


def test_extract_rst_text_empty_returns_placeholder(tmp_path):
    rst_file = tmp_path / "empty.rst"
    rst_file.write_text("", encoding="utf-8")
    result = ingest.extract_rst_text(rst_file)
    assert result == "[RST parsed but no extractable text found: empty.rst]"


def test_extract_text_dispatches_rst_extension(tmp_path):
    rst_file = tmp_path / "page.rst"
    rst_file.write_text(
        "Hello\n=====\n\nWorld.\n",
        encoding="utf-8",
    )
    result = ingest.extract_text(rst_file)
    assert "Hello" in result
    assert "World" in result


def test_extract_rst_text_latin1_fallback(tmp_path):
    rst_file = tmp_path / "latin1.rst"
    # Write bytes that are valid Latin-1 but invalid UTF-8
    rst_file.write_bytes(b"Title\n=====\n\ncaf\xe9\n")
    result = ingest.extract_rst_text(rst_file)
    assert "caf" in result


def test_extract_rst_text_missing_docutils(tmp_path, monkeypatch):
    rst_file = tmp_path / "nodeps.rst"
    rst_file.write_text("Hello\n=====\n", encoding="utf-8")

    # Clear any cached import of docutils.core so the function re-imports it
    import sys as _sys
    saved_modules = {}
    for mod_name in list(_sys.modules):
        if mod_name.startswith("docutils"):
            saved_modules[mod_name] = _sys.modules.pop(mod_name)

    import builtins
    real_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "docutils.core" or name == "docutils":
            raise ImportError("No module named 'docutils'")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)
    try:
        result = ingest.extract_rst_text(rst_file)
        assert result == "[RST support requires docutils: pip install docutils]"
    finally:
        # Restore docutils modules so other tests are not affected
        _sys.modules.update(saved_modules)


# ---------- DOCX extraction tests ----------

def test_extract_docx_text_multiple_paragraphs(tmp_path):
    """Integration test: DOCX with multiple paragraphs -> text extracted."""
    from docx import Document as DocxDocument
    docx_file = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph of the document.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.add_paragraph("Third and final paragraph.")
    doc.save(str(docx_file))

    result = ingest.extract_docx_text(docx_file)
    assert "First paragraph of the document." in result
    assert "Second paragraph with more content." in result
    assert "Third and final paragraph." in result
    # Paragraphs should be joined with newlines
    assert "\n" in result


def test_extract_docx_text_empty_returns_placeholder(tmp_path):
    """Integration test: empty DOCX -> placeholder returned."""
    from docx import Document as DocxDocument
    docx_file = tmp_path / "empty.docx"
    doc = DocxDocument()
    doc.save(str(docx_file))

    result = ingest.extract_docx_text(docx_file)
    assert result == "[DOCX parsed but no extractable text found: empty.docx]"


def test_extract_text_dispatches_docx_extension(tmp_path):
    """extract_text() dispatches .docx to extract_docx_text()."""
    from docx import Document as DocxDocument
    docx_file = tmp_path / "routed.docx"
    doc = DocxDocument()
    doc.add_paragraph("Dispatch test content.")
    doc.save(str(docx_file))

    result = ingest.extract_text(docx_file)
    assert "Dispatch test content." in result


def test_extract_docx_text_missing_python_docx(tmp_path, monkeypatch):
    """When python-docx is not installed, returns graceful message."""
    from docx import Document as DocxDocument
    docx_file = tmp_path / "nodeps.docx"
    doc = DocxDocument()
    doc.add_paragraph("Content here.")
    doc.save(str(docx_file))

    # Remove cached docx modules so the function re-imports
    import sys as _sys
    saved_modules = {}
    for mod_name in list(_sys.modules):
        if mod_name.startswith("docx"):
            saved_modules[mod_name] = _sys.modules.pop(mod_name)

    import builtins
    real_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("No module named 'docx'")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)
    try:
        result = ingest.extract_docx_text(docx_file)
        assert result == "[DOCX support requires python-docx: pip install python-docx]"
    finally:
        _sys.modules.update(saved_modules)


def test_extract_docx_text_whitespace_only_paragraphs(tmp_path):
    """DOCX with only whitespace paragraphs returns placeholder."""
    from docx import Document as DocxDocument
    docx_file = tmp_path / "whitespace.docx"
    doc = DocxDocument()
    doc.add_paragraph("   ")
    doc.add_paragraph("\t")
    doc.add_paragraph("")
    doc.save(str(docx_file))

    result = ingest.extract_docx_text(docx_file)
    assert result == "[DOCX parsed but no extractable text found: whitespace.docx]"
